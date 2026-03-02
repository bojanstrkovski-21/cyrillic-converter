#!/usr/bin/env python3
"""
Macedonian Font Converter
Converts old Macedonian font encoding (Latin characters) to proper Macedonian Cyrillic
Supports: .doc, .docx, .xls, .xlsx files
"""

import sys
import os
from pathlib import Path

# Character mapping from old font to Macedonian Cyrillic
CONVERSION_MAP = {
        'A': 'А', 'B': 'Б', 'C': 'Ц', 'D': 'Д', 'E': 'Е', 'F': 'Ф',
        'G': 'Г', 'H': 'Х', 'I': 'И', 'J': 'Ј', 'K': 'К', 'L': 'Л',
        'M': 'М', 'N': 'Н', 'O': 'О', 'P': 'П', 'Q': 'Љ', 'R': 'Р',
        'S': 'С', 'T': 'Т', 'U': 'У', 'V': 'В', 'W': 'Њ', 'X': 'Џ',
        'Y': 'Ѕ', 'Z': 'З',
        'a': 'а', 'b': 'б', 'c': 'ц', 'd': 'д', 'e': 'е', 'f': 'ф',
        'g': 'г', 'h': 'х', 'i': 'и', 'j': 'ј', 'k': 'к', 'l': 'л',
        'm': 'м', 'n': 'н', 'o': 'о', 'p': 'п', 'q': 'љ', 'r': 'р',
        's': 'с', 't': 'т', 'u': 'у', 'v': 'в', 'w': 'њ', 'x': 'џ',
        'y': 'ѕ', 'z': 'з',
        '[': 'ш', ']': 'ѓ', '\\': 'ж', ';': 'ч', ':': 'Ч',
        '{': 'Ш', '}': 'Ѓ', '|': 'Ж', '"': 'Ќ', '<': ';', '>': ':',
        '"': 'Ќ', '"': 'Ќ', '\'': 'ќ'
}

def convert_text(text):
    """Convert text using the character mapping"""
    if not text:
        return text
    
    result = []
    for char in text:
        result.append(CONVERSION_MAP.get(char, char))
    return ''.join(result)


def convert_docx(input_path, output_path):
    """Convert .docx file"""
    try:
        from docx import Document
    except ImportError:
        print("Installing required package: python-docx")
        import subprocess
        subprocess.check_call([sys.executable, '-m', 'pip', 'install', 'python-docx', '--break-system-packages', '-q'])
        from docx import Document
    
    doc = Document(input_path)
    
    # Convert paragraphs
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.text:
                run.text = convert_text(run.text)
    
    # Convert tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if run.text:
                            run.text = convert_text(run.text)
    
    doc.save(output_path)
    print(f"✓ Converted: {output_path}")


def convert_doc(input_path, output_path):
    """Convert .doc file (old Word format)"""
    try:
        import win32com.client
    except ImportError:
        print("Error: .doc file conversion requires Microsoft Word on Windows")
        print("Please convert your .doc files to .docx first, or run this on Windows with Word installed")
        return
    
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    
    try:
        doc = word.Documents.Open(str(Path(input_path).absolute()))
        
        # Convert main document text
        for paragraph in doc.Paragraphs:
            text = paragraph.Range.Text
            if text:
                paragraph.Range.Text = convert_text(text)
        
        # Save as .docx
        doc.SaveAs2(str(Path(output_path).absolute()), FileFormat=16)  # 16 = docx
        doc.Close()
        print(f"✓ Converted: {output_path}")
    finally:
        word.Quit()


def convert_xlsx(input_path, output_path):
    """Convert .xlsx file"""
    try:
        from openpyxl import load_workbook
    except ImportError:
        print("Installing required package: openpyxl")
        import subprocess
        subprocess.check_call([sys.executable, '-m', 'pip', 'install', 'openpyxl', '--break-system-packages', '-q'])
        from openpyxl import load_workbook
    
    wb = load_workbook(input_path)
    
    for sheet in wb.worksheets:
        for row in sheet.iter_rows():
            for cell in row:
                if cell.value and isinstance(cell.value, str):
                    cell.value = convert_text(cell.value)
    
    wb.save(output_path)
    print(f"✓ Converted: {output_path}")


def convert_xls(input_path, output_path):
    """Convert .xls file (old Excel format)"""
    try:
        import xlrd
        from openpyxl import Workbook
    except ImportError:
        print("Installing required packages: xlrd, openpyxl")
        import subprocess
        subprocess.check_call([sys.executable, '-m', 'pip', 'install', 'xlrd', 'openpyxl', '--break-system-packages', '-q'])
        import xlrd
        from openpyxl import Workbook
    
    # Read old format
    old_wb = xlrd.open_workbook(input_path)
    
    # Create new workbook
    new_wb = Workbook()
    new_wb.remove(new_wb.active)  # Remove default sheet
    
    for sheet_idx in range(old_wb.nsheets):
        old_sheet = old_wb.sheet_by_index(sheet_idx)
        new_sheet = new_wb.create_sheet(title=old_sheet.name)
        
        for row_idx in range(old_sheet.nrows):
            for col_idx in range(old_sheet.ncols):
                cell_value = old_sheet.cell_value(row_idx, col_idx)
                if isinstance(cell_value, str):
                    cell_value = convert_text(cell_value)
                new_sheet.cell(row=row_idx + 1, column=col_idx + 1, value=cell_value)
    
    new_wb.save(output_path)
    print(f"✓ Converted: {output_path}")


def convert_file(input_path):
    """Convert a single file based on its extension"""
    input_path = Path(input_path)
    
    if not input_path.exists():
        print(f"✗ Error: File not found: {input_path}")
        return False
    
    # Create output filename
    output_path = input_path.parent / f"{input_path.stem}_converted{input_path.suffix}"
    
    ext = input_path.suffix.lower()
    
    try:
        if ext == '.docx':
            convert_docx(str(input_path), str(output_path))
        elif ext == '.doc':
            convert_doc(str(input_path), str(output_path))
        elif ext == '.xlsx':
            convert_xlsx(str(input_path), str(output_path))
        elif ext == '.xls':
            convert_xls(str(input_path), str(output_path))
        else:
            print(f"✗ Unsupported file type: {ext}")
            return False
        
        return True
    except Exception as e:
        print(f"✗ Error converting {input_path}: {e}")
        return False


def main():
    if len(sys.argv) < 2:
        print("Macedonian Font Converter")
        print("=" * 50)
        print("\nUsage: python macedonian_converter.py <file1> [file2] [file3] ...")
        print("\nSupported formats: .doc, .docx, .xls, .xlsx")
        print("\nExample: python macedonian_converter.py document.docx spreadsheet.xlsx")
        print("\nConverted files will be saved with '_converted' suffix")
        sys.exit(1)
    
    print("Macedonian Font Converter")
    print("=" * 50)
    print()
    
    files = sys.argv[1:]
    success_count = 0
    
    for file_path in files:
        if convert_file(file_path):
            success_count += 1
    
    print()
    print(f"Conversion complete: {success_count}/{len(files)} files converted successfully")


if __name__ == "__main__":
    main()
